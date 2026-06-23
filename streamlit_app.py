"""
股票公告智能分析 Web 界面
功能：输入股票代码 → 显示该公司最近公告 + AI 分析结果

运行方式:
    streamlit run app.py
"""

import streamlit as st
import requests
import sqlite3
import pandas as pd
import numpy as np
import re
from datetime import datetime, timedelta
import os
import json
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# K线与财报可视化
import baostock as bs
import plotly.graph_objects as go
from plotly.subplots import make_subplots

# 云端/本地 统一配置（路径全由环境变量控制，默认保持本地行为）
from cloud_config import (
    get_db_connection,
    get_db_path,
    OUTPUT_DIR,
    REPORT_BASE_URL,
    resolve_output_path,
    ensure_dir,
)

# ============================
# 页面配置
# ============================
st.set_page_config(
    page_title="股票公告智能分析",
    page_icon="📈",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown("""
<style>
.bullish { color: #d32f2f; font-weight: bold; }
.bearish { color: #388e3c; font-weight: bold; }
.neutral { color: #f57c00; font-weight: bold; }
.tag-pill {
    display: inline-block;
    padding: 2px 10px;
    border-radius: 12px;
    font-size: 0.85em;
    background-color: #e3f2fd;
    color: #1565c0;
    margin-right: 6px;
}
</style>
""", unsafe_allow_html=True)

# ============================
# 1. 规则分类引擎（复用 announce_ai_processor.py）
# ============================
RULES = [
    ("业绩预告", ["业绩预告", "业绩快报", "一季度报告", "三季度报告", "半年度报告", "年度报告全文"], "中性", 10),
    ("股东减持", ["减持计划", "减持进展", "减持结果", "减持股份", "拟减持"], "利空", 10),
    ("股东增持", ["增持计划", "增持进展", "增持结果", "增持股份", "拟增持"], "利好", 10),
    ("重大合同", ["中标", "中标公告", "签订协议", "合作协议", "采购合同", "销售合同", "框架协议", "战略合作协议", "联合实验室"], "利好", 10),
    ("关联交易", ["关联交易", "关联资金", "关联方", "关联担保"], "中性", 10),
    ("股份回购", ["回购股份", "回购注销", "回购进展", "回购方案"], "利好", 10),
    ("股权激励", ["股权激励", "限制性股票", "股票期权", "激励计划"], "中性", 10),
    ("股权质押", ["解除质押", "质押式回购", "质押股份", "补充质押"], "中性", 10),
    ("定增/配股", ["非公开发行", "定向增发", "配股", "增发", "发行境外上市股份", "H股发行", "发行股票"], "中性", 10),
    ("并购重组", ["重大资产重组", "资产购买", "资产出售", "股权转让", "收购公司", "出售股票资产"], "中性", 10),
    ("停牌复牌", ["停牌公告", "复牌公告", "停牌进展"], "中性", 10),
    ("分红派息", ["权益分派", "利润分配", "分红派息", "派息", "股息", "分红方案"], "利好", 10),
    ("对外担保", ["对外担保", "提供担保", "担保额度"], "中性", 10),
    ("募集资金", ["募集资金", "募投项目", "募集资金专项"], "中性", 10),
    ("债券发行", ["公司债", "中期票据", "短期融资券", "科技创新债券", "发行债券"], "中性", 10),
    ("人事变动", ["董事长辞职", "总经理辞职", "高管辞职", "聘任", "任免", "人事变动", "离职", "辞职"], "中性", 10),
    ("诉讼仲裁", ["诉讼", "仲裁", "起诉", "被诉", "判决", "裁定"], "利空", 10),
    ("行政处罚", ["行政处罚", "监管函", "警示函", "立案调查", "纪律处分"], "利空", 10),
    ("退市风险", ["终止上市", "风险警示", "退市风险"], "利空", 10),
    ("股份解禁", ["限售股解禁", "限售股份", "解除限售"], "利空", 10),
    ("股东减持", ["减持"], "利空", 5),
    ("股东增持", ["增持"], "利好", 5),
    ("重大合同", ["合同"], "利好", 5),
    ("股份回购", ["回购"], "利好", 5),
    ("定增/配股", ["发行股票"], "中性", 5),
    ("并购重组", ["收购", "出售资产"], "中性", 5),
    ("停牌复牌", ["停牌", "复牌"], "中性", 5),
    ("分红派息", ["分红"], "利好", 5),
    ("对外担保", ["担保"], "中性", 5),
    ("债券发行", ["债券", "票据"], "中性", 5),
    ("人事变动", ["董事长", "总经理", "高管"], "中性", 5),
    ("股份解禁", ["解禁"], "利空", 5),
    ("行政处罚", ["处罚"], "利空", 5),
]


def rule_classify(title: str, content: str = "") -> tuple:
    """基于标题+正文前500字的规则分类"""
    text = (title + " " + content[:500]).lower()
    matched = []
    for tag, keywords, signal, priority in RULES:
        for kw in keywords:
            if kw.lower() in text:
                matched.append((tag, signal, priority))
                break
    if not matched:
        if "决议" in text or "会议" in text:
            return "董事会/监事会决议", "中性"
        if "调研" in text or "投资者关系" in text:
            return "投资者关系", "中性"
        if "法律意见" in text or "意见书" in text:
            return "法律意见书", "中性"
        if "章程" in text:
            return "公司章程", "中性"
        if "股东大会" in text or "股东会" in text:
            return "股东大会", "中性"
        if "监事会" in text:
            return "监事会决议", "中性"
        if "董事会" in text:
            return "董事会决议", "中性"
        return "其他", "中性"
    matched.sort(key=lambda x: x[2], reverse=True)
    return matched[0][0], matched[0][1]


# ============================
# 2. 数据获取
# ============================
@st.cache_data(ttl=300, show_spinner=False)
def fetch_announcements_from_api(stock_code: str, page_size: int = 50) -> pd.DataFrame:
    """从东方财富 API 获取公告列表"""
    url = "https://np-anotice-stock.eastmoney.com/api/security/ann"
    params = {
        "sr": "-1",
        "page_size": str(page_size),
        "page_index": "1",
        "ann_type": "A",
        "stock_list": stock_code,
        "f_node": "0",
        "s_node": "0",
    }
    try:
        resp = requests.get(url, params=params, headers={"User-Agent": "Mozilla/5.0"}, timeout=30)
        resp.encoding = "utf-8"
        data = resp.json()
        items = data.get("data", {}).get("list", [])
        if not items:
            return pd.DataFrame()

        rows = []
        for item in items:
            rows.append({
                "stock_code": stock_code,
                "stock_name": item.get("codes", [{}])[0].get("stock_name", ""),
                "title": item.get("title", ""),
                "category": item.get("column_code", ""),
                "announce_date": item.get("notice_date", "").split(" ")[0] if item.get("notice_date") else "",
                "url": f"https://data.eastmoney.com/notices/detail/{stock_code}/{item.get('art_code', '')}.html",
                "pdf_url": f"https://pdf.dfcfw.com/pdf/H2_{item.get('art_code', '')}_1.pdf",
                "art_code": item.get("art_code", ""),
            })
        return pd.DataFrame(rows)
    except Exception as e:
        st.error(f"从 API 获取公告失败: {e}")
        return pd.DataFrame()


def fetch_announcements_from_db(stock_code: str, db_path: str = None) -> pd.DataFrame:
    """从 SQLite 数据库获取公告（支持本地或云端同步后的数据库）"""
    try:
        conn = get_db_connection()
        df = pd.read_sql_query(
            """
            SELECT stock_code, stock_name, title, category, announce_date, url, pdf_url,
                   ai_tag, ai_summary, ai_signal, ai_reason
            FROM announcements
            WHERE stock_code = ?
            ORDER BY announce_date DESC
            """,
            conn,
            params=(stock_code,),
        )
        conn.close()
        return df
    except Exception:
        return pd.DataFrame()


def get_company_info(stock_code: str, db_path: str = None) -> dict:
    """从数据库获取公司基本信息（支持本地或云端同步后的数据库）"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute(
            "SELECT stock_name, industry, latest_price, total_market_cap FROM company_info WHERE stock_code = ?",
            (stock_code,),
        )
        row = cursor.fetchone()
        conn.close()
        if row:
            return {
                "stock_name": row[0] or "",
                "industry": row[1] or "",
                "latest_price": row[2],
                "total_market_cap": row[3],
            }
    except Exception:
        pass
    return {}


# ============================
# 2.5 baostock K线数据获取
# ============================
def convert_stock_code_for_baostock(stock_code: str) -> str:
    """将 6 位数字股票代码转换为 baostock 格式"""
    if stock_code.startswith(("60", "68", "51", "52", "53")):
        return f"sh.{stock_code}"
    elif stock_code.startswith(("00", "30", "15", "16", "18")):
        return f"sz.{stock_code}"
    elif stock_code.startswith(("43", "83", "87", "88", "82", "92")):
        return f"bj.{stock_code}"
    else:
        # 默认按上海处理
        return f"sh.{stock_code}"


@st.cache_data(ttl=3600, show_spinner=False)
def fetch_kline_from_baostock(stock_code: str, years: int = 8) -> pd.DataFrame:
    """从 baostock 获取近 N 年日线 K 线数据"""
    bs_code = convert_stock_code_for_baostock(stock_code)
    end_date = datetime.now().strftime("%Y-%m-%d")
    start_date = (datetime.now() - timedelta(days=years * 365)).strftime("%Y-%m-%d")

    try:
        lg = bs.login()
        if lg.error_code != "0":
            st.warning(f"baostock 登录失败: {lg.error_msg}")
            return pd.DataFrame()

        fields = "date,code,open,high,low,close,volume,amount,turn,pctChg,tradestatus"
        rs = bs.query_history_k_data_plus(
            bs_code,
            fields,
            start_date=start_date,
            end_date=end_date,
            frequency="d",
            adjustflag="3",  # 前复权
        )

        data_list = []
        while rs.next():
            data_list.append(rs.get_row_data())
        bs.logout()

        if not data_list:
            return pd.DataFrame()

        df = pd.DataFrame(data_list, columns=rs.fields)
        # 类型转换
        numeric_cols = ["open", "high", "low", "close", "volume", "amount", "turn", "pctChg"]
        for col in numeric_cols:
            if col in df.columns:
                df[col] = pd.to_numeric(df[col], errors="coerce")
        df["date"] = pd.to_datetime(df["date"])
        # 只保留交易日
        df = df[df.get("tradestatus", "1") == "1"].copy()
        return df
    except Exception as e:
        st.warning(f"获取 K 线数据失败: {e}")
        return pd.DataFrame()


def plot_kline(df_kline: pd.DataFrame, stock_name: str = "", days: int | None = None,
                flags: list[dict] | None = None, show_ma: bool = True,
                show_boll: bool = False, show_flags: bool = True) -> go.Figure:
    """使用 plotly 绘制 K 线图（含成交量），支持按天数缩放到最近范围，并标注旗形、均线、布林带"""
    if df_kline.empty:
        return go.Figure()

    df = df_kline.copy().reset_index(drop=True)
    df["date_str"] = df["date"].dt.strftime("%Y-%m-%d")
    df["ma5"] = df["close"].rolling(window=5).mean()
    df["ma10"] = df["close"].rolling(window=10).mean()
    df["ma20"] = df["close"].rolling(window=20).mean()
    df["ma60"] = df["close"].rolling(window=60).mean()

    # 布林带 BOLL(20,2)
    df["boll_mid"] = df["ma20"]
    df["boll_std"] = df["close"].rolling(window=20).std()
    df["boll_up"] = df["boll_mid"] + 2 * df["boll_std"]
    df["boll_down"] = df["boll_mid"] - 2 * df["boll_std"]

    x_vals = df["date_str"].tolist()

    fig = make_subplots(
        rows=2, cols=1,
        shared_xaxes=True,
        vertical_spacing=0.03,
        row_heights=[0.75, 0.25],
        subplot_titles=(f"{stock_name} K线走势", "成交量"),
    )

    # K线
    fig.add_trace(
        go.Candlestick(
            x=x_vals,
            open=df["open"],
            high=df["high"],
            low=df["low"],
            close=df["close"],
            name="K线",
            increasing_line_color="#d32f2f",
            decreasing_line_color="#388e3c",
            increasing_line_width=1,
            decreasing_line_width=1,
        ),
        row=1, col=1,
    )

    # 均线
    if show_ma:
        for col, color, name in [("ma5", "#F59E0B", "MA5"), ("ma10", "#06B6D4", "MA10"), ("ma20", "#3B82F6", "MA20"), ("ma60", "#8B5CF6", "MA60")]:
            fig.add_trace(
                go.Scatter(x=x_vals, y=df[col], mode="lines", name=name, line=dict(color=color, width=1)),
                row=1, col=1,
            )

    # 布林带
    if show_boll:
        fig.add_trace(
            go.Scatter(x=x_vals, y=df["boll_up"], mode="lines", name="BOLL上轨",
                       line=dict(color="rgba(128,128,128,0.6)", width=1, dash="dot")),
            row=1, col=1,
        )
        fig.add_trace(
            go.Scatter(x=x_vals, y=df["boll_mid"], mode="lines", name="BOLL中轨",
                       line=dict(color="rgba(128,128,128,0.8)", width=1.5)),
            row=1, col=1,
        )
        fig.add_trace(
            go.Scatter(x=x_vals, y=df["boll_down"], mode="lines", name="BOLL下轨",
                       line=dict(color="rgba(128,128,128,0.6)", width=1, dash="dot")),
            row=1, col=1,
        )

    # 成交量颜色
    colors = ["#d32f2f" if c >= o else "#388e3c" for c, o in zip(df["close"], df["open"])]
    fig.add_trace(
        go.Bar(x=x_vals, y=df["volume"], marker_color=colors, name="成交量", showlegend=False),
        row=2, col=1,
    )

    # 标注旗形（x 统一用字符串日期，与 category 轴完全匹配）
    if show_flags and flags:
        for idx, flag in enumerate(flags):
            is_bull = flag["type"] == "上飘旗"
            color = "rgba(211,47,47,0.15)" if is_bull else "rgba(56,142,60,0.15)"
            line_color = "#d32f2f" if is_bull else "#388e3c"

            ps = flag["pole_start"]
            pe = flag["pole_end"]
            fs = flag["flag_start"]
            fe = flag["flag_end"] - 1

            x_ps = x_vals[ps]
            x_pe = x_vals[pe]
            x_fs = x_vals[fs]
            x_fe = x_vals[fe]
            x_mid = x_vals[(fs + fe) // 2]

            # 旗杆矩形区域
            fig.add_shape(
                type="rect",
                x0=x_ps, x1=x_pe,
                y0=flag["pole_low"], y1=flag["pole_high"],
                fillcolor=color,
                line=dict(width=0),
                layer="below",
                row=1, col=1,
            )

            # 旗面上轨
            fig.add_shape(
                type="line",
                x0=x_fs, y0=flag["flag_high_start"],
                x1=x_fe, y1=flag["flag_high_end"],
                line=dict(color=line_color, width=2, dash="dash"),
                row=1, col=1,
            )

            # 旗面下轨
            fig.add_shape(
                type="line",
                x0=x_fs, y0=flag["flag_low_start"],
                x1=x_fe, y1=flag["flag_low_end"],
                line=dict(color=line_color, width=2, dash="dash"),
                row=1, col=1,
            )

            # 标注文字
            label_y = max(flag["flag_high_start"], flag["flag_high_end"]) * 1.01
            fig.add_annotation(
                x=x_mid, y=label_y,
                text=f"<b>{flag['type']}</b>",
                showarrow=False,
                font=dict(color=line_color, size=12),
                row=1, col=1,
            )

    fig.update_layout(
        height=600,
        xaxis_rangeslider_visible=False,
        hovermode="x unified",
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
        margin=dict(l=40, r=40, t=60, b=40),
        template="plotly_white",
    )
    fig.update_xaxes(
        type="category",
        tickmode="auto",
        nticks=10,
        showgrid=True, gridwidth=1, gridcolor="rgba(128,128,128,0.1)",
    )
    fig.update_yaxes(showgrid=True, gridwidth=1, gridcolor="rgba(128,128,128,0.1)")
    fig.update_yaxes(title_text="价格", row=1, col=1)
    fig.update_yaxes(title_text="成交量", row=2, col=1)

    # 缩放到指定天数范围（category 轴用索引）
    if days is not None:
        end_date = df["date"].max()
        start_date = end_date - pd.Timedelta(days=days)
        mask = df["date"] >= start_date
        if mask.any():
            start_idx = max(0, df[mask].index[0] - 1)
        else:
            start_idx = 0
        end_idx = len(df) - 1
        fig.update_xaxes(range=[start_idx, end_idx])
        # y 轴不固定，让 plotly 根据当前可见的 x 范围自动计算 y 范围
        fig.update_yaxes(autorange=True, row=1, col=1)

    return fig


def analyze_kline_tech(df_kline: pd.DataFrame) -> dict:
    """K线技术面分析：均线排列 + 成交量 + 量价关系"""
    if df_kline.empty or len(df_kline) < 20:
        return {}

    df = df_kline.copy()
    df["ma5"] = df["close"].rolling(window=5).mean()
    df["ma10"] = df["close"].rolling(window=10).mean()
    df["ma20"] = df["close"].rolling(window=20).mean()
    df["vol_ma5"] = df["volume"].rolling(window=5).mean()
    df["vol_ma20"] = df["volume"].rolling(window=20).mean()

    latest = df.iloc[-1]
    prev = df.iloc[-2] if len(df) > 1 else latest

    # 1. 均线排列判断
    ma5, ma10, ma20 = latest["ma5"], latest["ma10"], latest["ma20"]
    if pd.notna(ma5) and pd.notna(ma10) and pd.notna(ma20):
        if ma5 > ma10 > ma20:
            ma_trend = "多头排列"
            ma_desc = "短期均线在上，趋势偏强"
        elif ma5 < ma10 < ma20:
            ma_trend = "空头排列"
            ma_desc = "短期均线在下，趋势偏弱"
        else:
            ma_trend = "均线纠缠"
            ma_desc = "多空交织，方向不明"
    else:
        ma_trend = "数据不足"
        ma_desc = "均线数据不完整"

    # 2. 成交量判断
    vol = latest["volume"]
    vol_ma5 = latest["vol_ma5"]
    vol_ma20 = latest["vol_ma20"]
    if pd.notna(vol) and pd.notna(vol_ma5) and pd.notna(vol_ma20):
        avg_vol = (vol_ma5 + vol_ma20) / 2
        ratio = vol / avg_vol if avg_vol > 0 else 0
        if ratio >= 1.5:
            vol_status = "放量"
            vol_desc = "资金异动，关注方向"
        elif ratio >= 0.8:
            vol_status = "正常"
            vol_desc = "市场常态，无特殊信号"
        elif ratio >= 0.5:
            vol_status = "缩量"
            vol_desc = "交投清淡，观望情绪浓"
        else:
            vol_status = "地量"
            vol_desc = "极度低迷，可能变盘前兆"
        vol_ratio = f"{ratio:.2f}倍"
    else:
        vol_status = "未知"
        vol_desc = "成交量数据不完整"
        vol_ratio = "—"

    # 3. 量价关系（结合当日涨跌）
    price_change = latest["close"] - prev["close"]
    is_up = price_change >= 0
    vol_price = f"{vol_status}{'上涨' if is_up else '下跌'}"
    if vol_status == "放量":
        vp_desc = "资金积极介入，趋势可能延续" if is_up else "资金出逃，需警惕风险"
    elif vol_status == "缩量":
        vp_desc = "上涨动能减弱" if is_up else "抛压减轻，可能接近底部"
    elif vol_status == "地量":
        vp_desc = "变盘前兆，密切关注"
    else:
        vp_desc = "量价配合常态"

    return {
        "ma_trend": ma_trend,
        "ma_desc": ma_desc,
        "ma5": f"{ma5:.2f}" if pd.notna(ma5) else "—",
        "ma10": f"{ma10:.2f}" if pd.notna(ma10) else "—",
        "ma20": f"{ma20:.2f}" if pd.notna(ma20) else "—",
        "vol_status": vol_status,
        "vol_desc": vol_desc,
        "vol_ratio": vol_ratio,
        "vol": f"{vol/1e4:.0f}万" if pd.notna(vol) else "—",
        "vol_ma5": f"{vol_ma5/1e4:.0f}万" if pd.notna(vol_ma5) else "—",
        "vol_ma20": f"{vol_ma20/1e4:.0f}万" if pd.notna(vol_ma20) else "—",
        "vol_price": vol_price,
        "vp_desc": vp_desc,
        "price_change": price_change,
    }


# ============================
# 2.5 旗形检测
# ============================
def detect_flag_patterns(df_kline: pd.DataFrame) -> list[dict]:
    """
    检测上飘旗(Bull Flag)和下飘旗(Bear Flag)形态。
    返回检测到的旗形列表，每个元素包含类型、旗杆区间、旗面区间、画线坐标。
    """
    if len(df_kline) < 20:
        return []

    df = df_kline.copy().reset_index(drop=True)
    closes = df["close"].values
    highs = df["high"].values
    lows = df["low"].values
    volumes = df["volume"].values
    dates = df["date"].values
    n = len(df)

    flags = []

    def linear_slope(y_vals):
        """计算线性回归斜率"""
        x = np.arange(len(y_vals))
        if len(x) < 2:
            return 0.0
        return np.polyfit(x, y_vals, 1)[0]

    # 扫描窗口参数
    pole_min_days, pole_max_days = 3, 6
    flag_min_days, flag_max_days = 5, 15
    pole_threshold = 0.10  # 旗杆涨跌幅阈值 10%

    i = pole_max_days
    while i < n - flag_min_days:
        # 1. 检测旗杆
        pole_end = i
        pole_start_candidates = []
        for ps in range(pole_end - pole_max_days, max(0, pole_end - pole_min_days)):
            change = (closes[pole_end] - closes[ps]) / closes[ps]
            if abs(change) >= pole_threshold:
                pole_start_candidates.append((ps, change))

        if not pole_start_candidates:
            i += 1
            continue

        # 取变化最大的一段作为旗杆
        pole_start, pole_change = max(pole_start_candidates, key=lambda x: abs(x[1]))
        is_bull_pole = pole_change > 0

        # 2. 检测旗面
        flag_start = pole_end
        best_flag = None

        for flag_end in range(flag_start + flag_min_days, min(n, flag_start + flag_max_days + 1)):
            flag_highs = highs[flag_start:flag_end]
            flag_lows = lows[flag_start:flag_end]

            high_slope = linear_slope(flag_highs)
            low_slope = linear_slope(flag_lows)

            # 上飘旗：旗面必须向下倾斜（与主趋势反向）
            if is_bull_pole:
                # 高点和低点都应向下倾斜（略允许小正斜率容忍）
                if high_slope > 0 or low_slope > 0:
                    continue
                # 旗面回调幅度不能太大（不超过旗杆的50%）
                max_retrace = (highs[flag_start] - lows[flag_end - 1]) / (highs[flag_start] - lows[pole_start])
                if max_retrace > 0.6:
                    continue
                # 确认是平行通道（上轨和下轨斜率相近）
                slope_diff = abs(high_slope - low_slope)
                if slope_diff > abs(high_slope) * 1.5:
                    continue
                # 旗杆期间放量，旗面期间缩量
                pole_vol_avg = volumes[pole_start:pole_end].mean()
                flag_vol_avg = volumes[flag_start:flag_end].mean()
                if flag_vol_avg > pole_vol_avg * 0.8:
                    continue

                best_flag = {
                    "type": "上飘旗",
                    "pole_start": pole_start,
                    "pole_end": pole_end,
                    "flag_start": flag_start,
                    "flag_end": flag_end,
                    "pole_high": highs[pole_start:pole_end + 1].max(),
                    "pole_low": lows[pole_start:pole_end + 1].min(),
                    "flag_high_start": highs[flag_start],
                    "flag_high_end": highs[flag_end - 1],
                    "flag_low_start": lows[flag_start],
                    "flag_low_end": lows[flag_end - 1],
                    "date_pole_start": dates[pole_start],
                    "date_pole_end": dates[pole_end],
                    "date_flag_start": dates[flag_start],
                    "date_flag_end": dates[flag_end - 1],
                }
                break

            # 下飘旗：旗面必须向上倾斜（与主趋势反向）
            else:
                if high_slope < 0 or low_slope < 0:
                    continue
                # 旗面反弹幅度不能太大（不超过旗杆的50%）
                max_retrace = (highs[flag_end - 1] - lows[flag_start]) / (highs[pole_start] - lows[pole_start])
                if max_retrace > 0.6:
                    continue
                slope_diff = abs(high_slope - low_slope)
                if slope_diff > abs(high_slope) * 1.5:
                    continue
                pole_vol_avg = volumes[pole_start:pole_end].mean()
                flag_vol_avg = volumes[flag_start:flag_end].mean()
                if flag_vol_avg > pole_vol_avg * 0.8:
                    continue

                best_flag = {
                    "type": "下飘旗",
                    "pole_start": pole_start,
                    "pole_end": pole_end,
                    "flag_start": flag_start,
                    "flag_end": flag_end,
                    "pole_high": highs[pole_start:pole_end + 1].max(),
                    "pole_low": lows[pole_start:pole_end + 1].min(),
                    "flag_high_start": highs[flag_start],
                    "flag_high_end": highs[flag_end - 1],
                    "flag_low_start": lows[flag_start],
                    "flag_low_end": lows[flag_end - 1],
                    "date_pole_start": dates[pole_start],
                    "date_pole_end": dates[pole_end],
                    "date_flag_start": dates[flag_start],
                    "date_flag_end": dates[flag_end - 1],
                }
                break

        if best_flag:
            flags.append(best_flag)
            i = best_flag["flag_end"] + 1  # 跳过已检测到的旗形
        else:
            i += 1

    return flags


# ============================
# 2.6 财报数据读取
# ============================
FINANCIAL_KEYWORDS = [
    "业绩预告", "业绩快报", "季度报告", "年度报告", "半年度报告",
    "一季度报告", "三季度报告", "中报", "年报", "季报",
]


def is_financial_report(title: str) -> bool:
    """判断公告标题是否属于财报类"""
    t = str(title).lower()
    for kw in FINANCIAL_KEYWORDS:
        if kw in t:
            return True
    return False


@st.cache_data(ttl=3600, show_spinner=False)
def get_financial_indicators(stock_code: str) -> pd.DataFrame:
    """从 08_主要财务指标.csv 读取该公司关键财务指标"""
    csv_path = "04_原始数据备份/08_主要财务指标.csv"
    if not os.path.exists(csv_path):
        return pd.DataFrame()
    try:
        df = pd.read_csv(csv_path, encoding="utf-8-sig")
        # 数据格式：选项,指标,20260331,20251231,...（列为日期）
        # 只需返回全部数据，由展示部分处理
        return df
    except Exception:
        return pd.DataFrame()


# ============================
# 3. 结构化深度分析生成
# ============================
def call_llm_analyze(title: str, tag: str = "", signal: str = "") -> str | None:
    """调用 LLM API 生成结构化深度分析（需配置 API Key）"""
    api_key = os.environ.get("MOONSHOT_API_KEY") or os.environ.get("OPENAI_API_KEY")
    if not api_key:
        return None

    prompt = (
        f"你是一位专业的金融分析师。请对以下上市公司公告进行深度分析，"
        f"生成一段约200字的结构化摘要，严格按以下四部分组织，每部分前加对应标签：\n\n"
        f"【核心事件】提炼公告最核心的业务/资本运作事件\n"
        f"【影响主体】列出直接涉及的公司、股东、合作方等\n"
        f"【关键数字】提取金额、股份比例、时间期限等量化信息（若无则写'不涉及'）\n"
        f"【潜在影响】分析对股价、业绩、战略层面的短期与长期影响\n\n"
        f"公告标题：{title}\n"
        f"AI分类：{tag or '其他'}\n"
        f"信号：{signal or '中性'}"
    )

    try:
        # 优先尝试 Moonshot (Kimi) 的 OpenAI 兼容接口
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        }
        payload = {
            "model": "moonshot-v1-8k",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.3,
            "max_tokens": 512,
        }
        # 尝试 Moonshot endpoint
        resp = requests.post(
            "https://api.moonshot.cn/v1/chat/completions",
            headers=headers,
            json=payload,
            timeout=60,
        )
        if resp.status_code == 200:
            return resp.json()["choices"][0]["message"]["content"].strip()
    except Exception:
        pass

    try:
        # 回退到 OpenAI 兼容接口（通用 endpoint）
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        }
        payload = {
            "model": "gpt-3.5-turbo",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.3,
            "max_tokens": 512,
        }
        resp = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            json=payload,
            timeout=60,
        )
        if resp.status_code == 200:
            return resp.json()["choices"][0]["message"]["content"].strip()
    except Exception:
        pass

    return None


def rule_based_analysis(title: str, tag: str, signal: str, stock_name: str = "") -> str:
    """基于规则的模板化结构化摘要（无 LLM 时的兜底方案）"""
    name = stock_name or "公司"
    t = title.lower()

    # 尝试提取金额
    amount_match = re.search(r'(\d+(?:\.\d+)?)\s*[万亿]?元', title)
    amount = amount_match.group(0) if amount_match else "具体金额未披露"

    # 尝试提取股份比例
    ratio_match = re.search(r'(\d+(?:\.\d+)?)%', title)
    ratio = ratio_match.group(0) if ratio_match else ""

    templates = {
        "重大合同": (
            f"【核心事件】{name}签署/中标重大合同或战略合作协议，拓展业务版图。"
            f"【影响主体】{name}及合同相对方。"
            f"【关键数字】合同金额约{amount}。"
            f"【潜在影响】利好公司未来营收增长与市场份额提升，增强行业竞争力；"
            f"具体业绩贡献取决于合同执行进度与回款情况。"
        ),
        "股东减持": (
            f"【核心事件】{name}股东披露减持计划或减持进展，拟减持所持股份。"
            f"【影响主体】{name}、减持股东及二级市场投资者。"
            f"【关键数字】减持比例约{ratio or '待披露'}。"
            f"【潜在影响】股东减持通常被市场解读为利空信号，可能短期对股价形成压力；"
            f"但若减持用于自身资金需求且比例可控，长期影响有限。"
        ),
        "股东增持": (
            f"【核心事件】{name}股东或管理层计划增持公司股份，彰显对公司发展的信心。"
            f"【影响主体】{name}、增持方及全体股东。"
            f"【关键数字】增持比例约{ratio or '待披露'}。"
            f"【潜在影响】增持行为通常被市场视为利好信号，有助于稳定股价预期；"
            f"体现内部人对公司价值的认可，提振投资者信心。"
        ),
        "股份回购": (
            f"【核心事件】{name}推进股份回购计划，拟以自有资金回购公司股份。"
            f"【影响主体】{name}及全体股东。"
            f"【关键数字】回购金额/比例待披露。"
            f"【潜在影响】回购注销可减少总股本、提升每股收益，属于利好信号；"
            f"同时彰显公司现金流充裕、管理层认为股价被低估。"
        ),
        "分红派息": (
            f"【核心事件】{name}披露年度/半年度利润分配方案，拟向股东派发现金红利或送转股份。"
            f"【影响主体】{name}及全体股东。"
            f"【关键数字】分红金额/送转比例待披露。"
            f"【潜在影响】分红体现公司盈利能力和现金回报意愿，属利好信号；"
            f"高股息率对长期投资者具有吸引力，除权后需关注填权行情。"
        ),
        "业绩预告": (
            f"【核心事件】{name}发布业绩预告/快报，披露当期主要经营成果。"
            f"【影响主体】{name}、投资者及监管机构。"
            f"【关键数字】业绩增减幅度待披露。"
            f"【潜在影响】业绩向好为利好，不及预期为利空；"
            f"业绩预告是投资者判断公司基本面的重要窗口，关注同比环比变化。"
        ),
        "股权激励": (
            f"【核心事件】{name}实施股权激励计划，向核心员工授予限制性股票或期权。"
            f"【影响主体】{name}、激励对象及全体股东。"
            f"【关键数字】激励股份数量及业绩考核指标待披露。"
            f"【潜在影响】利好长期发展，有助于绑定核心人才、激发经营活力；"
            f"短期可能因股份支付费用对利润产生一定摊薄。"
        ),
        "定增/配股": (
            f"【核心事件】{name}推进非公开发行/定向增发/配股等再融资事项。"
            f"【影响主体】{name}、认购方及现有股东。"
            f"【关键数字】募资规模及发行价格待披露。"
            f"【潜在影响】中性偏利好：募资用于扩产或并购可打开成长空间；"
            f"但增发会稀释现有股东权益，需关注资金用途及定价折让。"
        ),
        "并购重组": (
            f"【核心事件】{name}推进资产购买、出售或股权重组事项。"
            f"【影响主体】{name}、交易对手方及标的资产。"
            f"【关键数字】交易对价及资产估值待披露。"
            f"【潜在影响】若并购优质资产则为利好，出售非核心资产可优化结构；"
            f"需关注交易定价合理性及后续整合风险。"
        ),
        "人事变动": (
            f"【核心事件】{name}发生高管/董事/监事人事变动。"
            f"【影响主体】{name}、新任/离任人员及管理层。"
            f"【关键数字】不涉及。"
            f"【潜在影响】常规人事变动为中性；"
            f"若核心创始人或关键高管离任，可能引发市场对公司战略连续性的担忧。"
        ),
        "诉讼仲裁": (
            f"【核心事件】{name}涉及诉讼或仲裁事项。"
            f"【影响主体】{name}及诉讼相对方。"
            f"【关键数字】涉案金额待披露。"
            f"【潜在影响】通常被市场视为利空，可能产生或有负债或声誉风险；"
            f"需关注诉讼进展及最终判决结果对财务报表的影响。"
        ),
        "行政处罚": (
            f"【核心事件】{name}收到监管部门的行政处罚、警示函或立案调查通知。"
            f"【影响主体】{name}、监管机构及相关责任人。"
            f"【关键数字】处罚金额待披露。"
            f"【潜在影响】利空信号，反映公司治理或合规存在瑕疵；"
            f"严重处罚可能影响融资、投标等经营活动，需关注后续整改情况。"
        ),
        "停牌复牌": (
            f"【核心事件】{name}因重大事项申请股票停牌/复牌。"
            f"【影响主体】{name}及全体投资者。"
            f"【关键数字】停牌期限待披露。"
            f"【潜在影响】停牌期间无法交易；"
            f"复牌后股价走势取决于停牌期间事项的实质性影响。"
        ),
        "对外担保": (
            f"【核心事件】{name}为子公司或关联方提供对外担保。"
            f"【影响主体】{name}、被担保方及债权人。"
            f"【关键数字】担保额度待披露。"
            f"【潜在影响】中性：适度担保属正常经营行为；"
            f"但若担保规模过大或对象为高风险主体，可能增加或有负债风险。"
        ),
        "股权质押": (
            f"【核心事件】{name}股东办理股份质押或解除质押。"
            f"【影响主体】{name}、质押股东及质权方。"
            f"【关键数字】质押比例待披露。"
            f"【潜在影响】解除质押为利好，降低平仓风险；"
            f"新增高比例质押需警惕股价下跌引发的强制平仓风险。"
        ),
        "股份解禁": (
            f"【核心事件】{name}限售股解禁，相关股份可在二级市场流通。"
            f"【影响主体】{name}、限售股股东及二级市场投资者。"
            f"【关键数字】解禁股份数量及占总股本比例待披露。"
            f"【潜在影响】通常被市场视为利空，解禁后股东减持可能增加抛压；"
            f"但若股东承诺不减持或公司基本面强劲，影响可对冲。"
        ),
        "公司章程": (
            f"【核心事件】{name}修订或披露公司章程。"
            f"【影响主体】{name}及全体股东。"
            f"【关键数字】不涉及。"
            f"【潜在影响】中性：属常规公司治理行为；"
            f"若涉及特别表决权、反收购条款等重大修改，需关注对股东权益的影响。"
        ),
        "关联交易": (
            f"【核心事件】{name}发生关联交易事项。"
            f"【影响主体】{name}及关联方。"
            f"【关键数字】交易金额待披露。"
            f"【潜在影响】中性：合规披露的关联交易属正常经营；"
            f"需关注交易定价是否公允，是否存在利益输送风险。"
        ),
        "募集资金": (
            f"【核心事件】{name}披露募集资金使用/变更/结项情况。"
            f"【影响主体】{name}及投资者。"
            f"【关键数字】募集资金金额及使用进度待披露。"
            f"【潜在影响】中性：募投项目顺利推进利好长期发展；"
            f"若频繁变更用途或进度滞后，需关注管理层执行能力。"
        ),
        "债券发行": (
            f"【核心事件】{name}拟发行公司债、中期票据或其他债券融资工具。"
            f"【影响主体】{name}及债券投资者。"
            f"【关键数字】发行规模及票面利率待披露。"
            f"【潜在影响】中性：债券融资可优化债务结构、补充流动资金；"
            f"需关注公司偿债能力及债券评级情况。"
        ),
    }

    if tag in templates:
        return templates[tag]

    # 兜底模板
    return (
        f"【核心事件】{name}发布公告：{title}。"
        f"【影响主体】{name}及相关利益方。"
        f"【关键数字】具体数据请查阅公告原文。"
        f"【潜在影响】该公告被AI规则分类为「{tag}」，信号「{signal}」；"
        f"具体影响需结合公告全文及公司基本面综合判断。"
    )


def get_structured_summary(row: pd.Series, stock_name: str = "") -> str:
    """获取结构化摘要：优先用数据库 → LLM → 规则模板"""
    # 1. 本地数据库已有结构化摘要且包含四个标签
    existing = str(row.get("ai_summary", ""))
    if existing and existing != "nan" and all(tag in existing for tag in ["核心事件", "影响主体", "关键数字", "潜在影响"]):
        return existing

    # 2. 尝试调用 LLM（需配置 API Key）
    llm_result = call_llm_analyze(row.get("title", ""), row.get("ai_tag", ""), row.get("ai_signal", ""))
    if llm_result:
        return llm_result

    # 3. 规则模板兜底
    return rule_based_analysis(
        row.get("title", ""),
        row.get("ai_tag", "其他"),
        row.get("ai_signal", "中性"),
        stock_name,
    )


def get_reason(row: pd.Series) -> str:
    """获取判断理由：优先用数据库 → 规则模板"""
    existing = str(row.get("ai_reason", ""))
    if existing and existing != "nan":
        return existing

    # 基于标签的简要判断理由
    tag = row.get("ai_tag", "其他")
    signal = row.get("ai_signal", "中性")
    title = row.get("title", "")

    reasons = {
        "重大合同": f"重大合同/战略合作通常被市场解读为利好，意味着公司获得增量订单或技术合作，对未来营收有积极影响。",
        "股东减持": f"股东减持计划通常被市场视为利空信号，短期内可能对股价形成压力；需关注减持比例及实际执行情况。",
        "股东增持": f"股东/管理层增持彰显对公司长期发展的信心，通常被视为利好信号，有助于稳定市场预期。",
        "股份回购": f"股份回购可减少总股本、提升每股收益，彰显现金流充裕和管理层对股价的认可，属利好信号。",
        "分红派息": f"分红派息体现公司盈利能力和现金回报意愿，高股息对长期投资者有吸引力，通常被市场视为利好。",
        "业绩预告": f"业绩预告是投资者判断基本面的重要窗口，业绩向好为利好，不及预期为利空。",
        "股权激励": f"股权激励有助于绑定核心人才、激发经营活力，利好公司长期发展，但短期可能摊薄利润。",
        "定增/配股": f"再融资事项中性偏利好：若用于扩产/并购可打开成长空间，但会稀释现有股东权益。",
        "并购重组": f"并购重组的影响取决于标的资产质量和交易定价，优质资产注入为利好，出售非核心资产可优化结构。",
        "人事变动": f"常规人事变动为中性；若核心高管离任，可能引发市场对公司战略连续性的担忧。",
        "诉讼仲裁": f"涉诉事项通常被市场视为利空，可能产生或有负债或声誉风险，需关注诉讼进展及最终判决。",
        "行政处罚": f"监管处罚为利空信号，反映公司治理或合规存在瑕疵，严重处罚可能影响后续融资及经营活动。",
        "停牌复牌": f"停牌期间无法交易，复牌后股价走势取决于停牌期间事项的实质性影响。",
        "对外担保": f"适度担保属正常经营行为，但若担保规模过大或对象为高风险主体，可能增加或有负债风险。",
        "股权质押": f"解除质押为利好（降低平仓风险），新增高比例质押需警惕股价下跌引发的强制平仓风险。",
        "股份解禁": f"限售股解禁通常被市场视为利空，解禁后股东减持可能增加抛压，需关注股东承诺。",
        "公司章程": f"章程修订属常规公司治理行为，但若涉及特别表决权等重大修改，需关注对股东权益的影响。",
        "关联交易": f"合规披露的关联交易属正常经营，需关注交易定价是否公允，是否存在利益输送风险。",
        "募集资金": f"募投项目顺利推进利好长期发展，若频繁变更用途或进度滞后，需关注管理层执行能力。",
        "债券发行": f"债券融资可优化债务结构、补充流动资金，需关注公司偿债能力及债券评级情况。",
        "董事会决议": f"董事会决议本身为程序性事件，核心内容已在单独公告中披露，需关注后续股东会审议结果。",
        "监事会决议": f"监事会决议为常规公司治理程序，属中性事件。",
        "股东大会": f"股东大会决议反映股东意志，对公司重大决策有约束力，需关注投票结果及反对票比例。",
        "投资者关系": f"投资者关系活动为中性事件，有助于增进市场对公司的了解。",
        "法律意见书": f"法律意见书为合规性文件，属中性信息披露。",
        "退市风险": f"退市风险警示为重大利空信号，投资者需高度关注公司基本面改善情况。",
    }

    if tag in reasons:
        return reasons[tag]
    return f"该公告被AI规则分类为「{tag}」，信号「{signal}」；具体影响需结合公告全文及公司基本面综合判断。"


def generate_report_md(stock_code: str, stock_name: str, row: pd.Series) -> str:
    """按照 02_AI深度分析报告 模板生成单条公告深度分析报告（Markdown格式）"""
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    date = row.get("announce_date", "")
    title = row.get("title", "")
    tag = row.get("ai_tag", "其他")
    signal = row.get("ai_signal", "中性")
    summary = get_structured_summary(row, stock_name)
    reason = get_reason(row)

    md = f"""# {stock_name}({stock_code}) 公告AI深度分析报告

> 分析时间: {now}
> 分析引擎: Kimi AI + 规则分类引擎

---

## {date}｜{title}

- **AI分类**: {tag}
- **信号判断**: {signal}
- **结构化摘要**: {summary}
- **判断理由**: {reason}

---
"""
    return md


def generate_report_docx(stock_code: str, stock_name: str, row: pd.Series) -> bytes:
    """按照 02_AI深度分析报告 模板生成单条公告深度分析报告（Word .docx格式）"""
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    date = row.get("announce_date", "")
    title = row.get("title", "")
    tag = row.get("ai_tag", "其他")
    signal = row.get("ai_signal", "中性")
    summary = get_structured_summary(row, stock_name)
    reason = get_reason(row)

    doc = Document()

    # 标题
    heading = doc.add_heading(f"{stock_name}({stock_code}) 公告AI深度分析报告", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 元信息
    meta = doc.add_paragraph()
    meta.add_run(f"分析时间: {now}\n").italic = True
    meta.add_run("分析引擎: Kimi AI + 规则分类引擎").italic = True

    # 分隔线
    doc.add_paragraph("─" * 40)

    # 公告标题（二级标题）
    doc.add_heading(f"{date}｜{title}", level=2)

    # AI分类
    p = doc.add_paragraph()
    p.add_run("AI分类: ").bold = True
    p.add_run(tag)

    # 信号判断
    p = doc.add_paragraph()
    p.add_run("信号判断: ").bold = True
    run = p.add_run(signal)
    if signal == "利好":
        run.font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)
    elif signal == "利空":
        run.font.color.rgb = RGBColor(0x38, 0x8E, 0x3C)
    else:
        run.font.color.rgb = RGBColor(0xF5, 0x7C, 0x00)
    run.bold = True

    # 结构化摘要
    p = doc.add_paragraph()
    p.add_run("结构化摘要: ").bold = True
    p.add_run(summary)

    # 判断理由
    p = doc.add_paragraph()
    p.add_run("判断理由: ").bold = True
    p.add_run(reason)

    # 分隔线
    doc.add_paragraph("─" * 40)

    # 保存到内存
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============================
# 4. 信号样式
# ============================
def signal_badge(signal: str) -> str:
    if signal == "利好":
        return '<span class="bullish">📈 利好</span>'
    elif signal == "利空":
        return '<span class="bearish">📉 利空</span>'
    else:
        return '<span class="neutral">➖ 中性</span>'


def tag_badge(tag: str) -> str:
    return f'<span class="tag-pill">{tag}</span>'


# ============================
# 4. 主界面
# ============================
st.title("📈 股票公告智能分析系统")
st.caption("输入股票代码，自动获取最新公告并进行 AI 规则分类分析")

# --- 输入区 ---
col1, col2 = st.columns([1, 4])
with col1:
    stock_code = st.text_input("股票代码", value="002340", max_chars=6)
with col2:
    st.write("")
    st.write("")
    search_clicked = st.button("🔍 查询公告", type="primary", use_container_width=False)

if not stock_code or not re.match(r"^\d{6}$", stock_code.strip()):
    st.warning("请输入 6 位数字股票代码（如 002340、600031）")
    st.stop()

stock_code = stock_code.strip()

# --- 数据加载 ---
if search_clicked or True:
    with st.spinner(f"正在加载 {stock_code} 的公告数据..."):
        # 先查本地数据库（含 AI 深度分析结果）
        df_local = fetch_announcements_from_db(stock_code)

        # 再查 API（获取最新数据）
        df_api = fetch_announcements_from_api(stock_code, page_size=50)

        if df_api.empty and df_local.empty:
            st.error(f"未找到股票 {stock_code} 的公告数据，请检查代码是否正确。")
            st.stop()

        # 合并：API 数据做基础，本地数据补充 AI 分析字段
        if not df_api.empty and not df_local.empty:
            df_api = df_api.merge(
                df_local[["title", "ai_tag", "ai_summary", "ai_signal", "ai_reason"]].drop_duplicates("title"),
                on="title",
                how="left",
            )
            df = df_api
        elif not df_api.empty:
            df = df_api
        else:
            df = df_local

        # 对缺少 AI 标签的行进行规则分类
        for idx, row in df.iterrows():
            if pd.isna(row.get("ai_tag")) or row.get("ai_tag") == "":
                tag, signal = rule_classify(row["title"], row.get("content_preview", ""))
                df.at[idx, "ai_tag"] = tag
                df.at[idx, "ai_signal"] = signal

        df = df.sort_values("announce_date", ascending=False).reset_index(drop=True)

    # --- 公司信息卡片 ---
    company = get_company_info(stock_code)
    name = company.get("stock_name") or (df["stock_name"].iloc[0] if not df.empty else "")

    info_cols = st.columns(4)
    with info_cols[0]:
        st.metric("股票代码", stock_code)
    with info_cols[1]:
        st.metric("公司名称", name or "—")
    with info_cols[2]:
        st.metric("所属行业", company.get("industry") or "—")
    with info_cols[3]:
        cap = company.get("total_market_cap")
        st.metric("总市值", f"{cap:,.0f} 亿" if cap else "—")

    st.divider()

    # --- K 线走势图 ---
    st.subheader("📈 近8年 K 线走势（前复权）")
    with st.spinner("正在从 baostock 获取 K 线数据..."):
        df_kline = fetch_kline_from_baostock(stock_code, years=8)
    if not df_kline.empty:
        # 初始化 session state（默认显示近3年）
        if "kline_days" not in st.session_state:
            st.session_state.kline_days = 1095

        # 初始化显示控制 session state
        if "show_ma" not in st.session_state:
            st.session_state.show_ma = True
        if "show_boll" not in st.session_state:
            st.session_state.show_boll = False
        if "show_flags" not in st.session_state:
            st.session_state.show_flags = False

        # 时间范围按钮（紧凑排列）
        st.markdown("<span style='font-size:12px;color:#666'>时间范围</span>", unsafe_allow_html=True)
        btn_cols = st.columns([0.5, 0.5, 0.5, 0.5, 0.5, 0.5, 0.5, 0.5, 10])
        ranges = [
            (btn_cols[0], "5日", 5),
            (btn_cols[1], "10日", 10),
            (btn_cols[2], "30日", 30),
            (btn_cols[3], "90日", 90),
            (btn_cols[4], "1年", 365),
            (btn_cols[5], "3年", 1095),
            (btn_cols[6], "5年", 1825),
            (btn_cols[7], "全部", None),
        ]
        for col, label, days in ranges:
            with col:
                is_active = st.session_state.kline_days == days
                if st.button(label, key=f"kline_{stock_code}_{days}", type="primary" if is_active else "secondary"):
                    st.session_state.kline_days = days
                    st.rerun()

        # 显示控制按钮
        st.markdown("<span style='font-size:12px;color:#666'>显示图层</span>", unsafe_allow_html=True)
        ctrl_cols = st.columns([0.7, 0.7, 0.7, 10])
        ctrl_items = [
            (ctrl_cols[0], "均线", "show_ma"),
            (ctrl_cols[1], "布林带", "show_boll"),
            (ctrl_cols[2], "旗形", "show_flags"),
        ]
        for col, label, state_key in ctrl_items:
            with col:
                is_on = st.session_state[state_key]
                if st.button(label, key=f"ctrl_{stock_code}_{state_key}", type="primary" if is_on else "secondary"):
                    st.session_state[state_key] = not is_on
                    st.rerun()

        # 旗形检测（仅在开启时计算）
        flags = detect_flag_patterns(df_kline) if st.session_state.show_flags else []

        fig_kline = plot_kline(
            df_kline,
            stock_name=name or stock_code,
            days=st.session_state.kline_days,
            flags=flags,
            show_ma=st.session_state.show_ma,
            show_boll=st.session_state.show_boll,
            show_flags=st.session_state.show_flags,
        )
        st.plotly_chart(fig_kline, use_container_width=True)

        # 旗形检测结果展示（仅在开启时显示）
        if st.session_state.show_flags and flags:
            st.markdown("<span style='font-size:13px'>**🚩 旗形形态检测**</span>", unsafe_allow_html=True)
            for f in flags:
                d_ps = pd.to_datetime(f["date_pole_start"]).strftime("%Y-%m-%d")
                d_pe = pd.to_datetime(f["date_pole_end"]).strftime("%Y-%m-%d")
                d_fs = pd.to_datetime(f["date_flag_start"]).strftime("%Y-%m-%d")
                d_fe = pd.to_datetime(f["date_flag_end"]).strftime("%Y-%m-%d")
                if f["type"] == "上飘旗":
                    st.success(
                        f"检测到 **上飘旗**（看涨信号）："
                        f"旗杆期 {d_ps} → {d_pe} 急速拉升，"
                        f"随后 {d_fs} → {d_fe} 缩量向下整理，"
                        f"关注放量突破旗面上轨后的跟进机会。"
                    )
                else:
                    st.error(
                        f"检测到 **下飘旗**（看跌信号）："
                        f"旗杆期 {d_ps} → {d_pe} 急速杀跌，"
                        f"随后 {d_fs} → {d_fe} 缩量向上整理，"
                        f"警惕放量跌破旗面下轨后的持续下跌风险。"
                    )

        # 技术面分析
        tech = analyze_kline_tech(df_kline)
        if tech:
            st.markdown("<span style='font-size:13px'>**📐 技术面诊断**</span>", unsafe_allow_html=True)
            tech_cols = st.columns(3)
            with tech_cols[0]:
                if tech["ma_trend"] == "多头排列":
                    st.metric("均线排列", tech["ma_trend"], tech["ma_desc"], delta_color="inverse")
                elif tech["ma_trend"] == "空头排列":
                    st.metric("均线排列", tech["ma_trend"], tech["ma_desc"])
                else:
                    st.metric("均线排列", tech["ma_trend"], tech["ma_desc"])
                st.caption(f"MA5:{tech['ma5']}  MA10:{tech['ma10']}  MA20:{tech['ma20']}")
            with tech_cols[1]:
                vp_color = "normal"
                if "放量" in tech["vol_status"]:
                    vp_color = "inverse"
                st.metric("成交量", tech["vol_status"], f"{tech['vol_desc']}（{tech['vol_ratio']}）", delta_color=vp_color)
                st.caption(f"当日:{tech['vol']}  5日均:{tech['vol_ma5']}  20日均:{tech['vol_ma20']}")
            with tech_cols[2]:
                if "上涨" in tech["vol_price"]:
                    st.metric("量价关系", tech["vol_price"], tech["vp_desc"], delta_color="inverse")
                else:
                    st.metric("量价关系", tech["vol_price"], tech["vp_desc"])

        kline_cols = st.columns(4)
        latest = df_kline.iloc[-1]
        prev = df_kline.iloc[-2] if len(df_kline) > 1 else latest
        kline_cols[0].metric("最新收盘价", f"{latest['close']:.2f}", f"{(latest['close'] - prev['close']):.2f}")
        # 计算20日平均成交额
        amount_ma20 = df_kline["amount"].tail(20).mean()
        kline_cols[1].metric("成交额", f"{latest['amount']/1e8:.2f} 亿")
        kline_cols[1].caption(f"20日均: {amount_ma20/1e8:.2f} 亿")
        kline_cols[2].metric("换手率", f"{latest['turn']:.2f}%")
        kline_cols[3].metric("8年涨跌幅", f"{((latest['close'] / df_kline.iloc[0]['close'] - 1) * 100):.2f}%")
    else:
        st.info("未获取到 K 线数据，请检查网络或股票代码。")

    st.divider()

    # --- 统计概览 ---
    st.subheader("📊 公告信号统计")
    sig_counts = df["ai_signal"].value_counts().to_dict()
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("总公告数", len(df))
    c2.metric("利好", sig_counts.get("利好", 0), delta_color="inverse")
    c3.metric("利空", sig_counts.get("利空", 0), delta_color="inverse")
    c4.metric("中性", sig_counts.get("中性", 0))

    # 分类分布
    tag_counts = df["ai_tag"].value_counts().head(10)
    chart_col, table_col = st.columns([1, 1])
    with chart_col:
        st.bar_chart(tag_counts, use_container_width=True, height=250)
    with table_col:
        st.dataframe(
            tag_counts.reset_index().rename(columns={"index": "分类", "ai_tag": "数量"}),
            use_container_width=True,
            hide_index=True,
        )

    st.divider()

    # --- 公告列表 ---
    st.subheader(f"📋 最近公告列表（共 {len(df)} 条）")

    # 筛选器
    filter_col1, filter_col2 = st.columns([1, 3])
    with filter_col1:
        signal_filter = st.selectbox("筛选信号", ["全部", "利好", "利空", "中性"])
    with filter_col2:
        search_title = st.text_input("搜索标题关键词", placeholder="输入关键词筛选公告...")

    filtered_df = df.copy()
    if signal_filter != "全部":
        filtered_df = filtered_df[filtered_df["ai_signal"] == signal_filter]
    if search_title:
        filtered_df = filtered_df[filtered_df["title"].str.contains(search_title, case=False, na=False)]

    # 展示表格
    display_df = filtered_df[["announce_date", "title", "ai_tag", "ai_signal"]].copy()
    display_df.columns = ["日期", "公告标题", "AI分类", "信号"]

    st.dataframe(
        display_df,
        use_container_width=True,
        hide_index=True,
        column_config={
            "日期": st.column_config.DateColumn("日期", format="YYYY-MM-DD"),
            "公告标题": st.column_config.TextColumn("公告标题", width="large"),
            "AI分类": st.column_config.TextColumn("AI分类", width="medium"),
            "信号": st.column_config.TextColumn("信号", width="small"),
        },
    )

    st.divider()

    # --- 财报公告专区 ---
    st.subheader("📑 财报公告专区")
    df_finance = df[df["title"].apply(is_financial_report)].copy()
    if not df_finance.empty:
        st.markdown(f"检测到 **{len(df_finance)}** 条财报相关公告，已单独列出：")
        fin_display = df_finance[["announce_date", "title", "ai_tag", "ai_signal"]].copy()
        fin_display.columns = ["日期", "公告标题", "AI分类", "信号"]
        st.dataframe(
            fin_display,
            use_container_width=True,
            hide_index=True,
            column_config={
                "日期": st.column_config.DateColumn("日期", format="YYYY-MM-DD"),
                "公告标题": st.column_config.TextColumn("公告标题", width="large"),
                "AI分类": st.column_config.TextColumn("AI分类", width="medium"),
                "信号": st.column_config.TextColumn("信号", width="small"),
            },
        )
    else:
        st.info("未检测到财报相关公告。")

    # --- 关键财务指标趋势 ---
    st.subheader("💰 关键财务指标趋势")
    df_fin_indicators = get_financial_indicators(stock_code)
    if not df_fin_indicators.empty:
        # 指标名称在第一列（'选项'）和第二列（'指标'），后续列为日期
        # 提取常用指标行
        indicators_of_interest = {
            "归母净利润": "归母净利润（元）",
            "营业总收入": "营业总收入（元）",
            "营业成本": "营业成本（元）",
            "净利润": "净利润（元）",
        }
        date_cols = [c for c in df_fin_indicators.columns if re.match(r"^\d{8}$", str(c))]
        date_cols_sorted = sorted(date_cols, reverse=True)[:16]  # 最近16期（约4年季度数据）
        date_cols_sorted = sorted(date_cols_sorted)  # 再按时间正序排

        fig_fin = go.Figure()
        colors = ["#d32f2f", "#1976d2", "#388e3c", "#f57c00"]
        color_idx = 0
        for ind_name, label in indicators_of_interest.items():
            row = df_fin_indicators[df_fin_indicators["指标"] == ind_name]
            if row.empty:
                continue
            values = []
            x_labels = []
            for d in date_cols_sorted:
                val = row.iloc[0].get(d)
                if pd.notna(val):
                    try:
                        values.append(float(val))
                        x_labels.append(f"{d[:4]}-{d[4:6]}")
                    except (ValueError, TypeError):
                        pass
            if values:
                # 亿元转换
                values_yi = [v / 1e8 for v in values]
                fig_fin.add_trace(
                    go.Scatter(
                        x=x_labels, y=values_yi, mode="lines+markers",
                        name=label, line=dict(color=colors[color_idx % len(colors)], width=2),
                    )
                )
                color_idx += 1

        if fig_fin.data:
            fig_fin.update_layout(
                height=400,
                xaxis_title="报告期",
                yaxis_title="金额（亿元）",
                hovermode="x unified",
                legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
                margin=dict(l=40, r=40, t=60, b=40),
                template="plotly_white",
            )
            st.plotly_chart(fig_fin, use_container_width=True)
        else:
            st.info("财务指标数据格式不匹配，无法绘制趋势图。")
    else:
        st.info("未找到主要财务指标数据（08_主要财务指标.csv）。")

    st.divider()

    # --- 逐条详情（可展开） ---
    st.subheader("🔍 公告详情与 AI 分析")
    for idx, row in filtered_df.head(20).iterrows():
        with st.expander(f"[{row['announce_date']}] {row['title']}"):
            cols = st.columns([1, 3])
            with cols[0]:
                st.markdown(f"**AI 分类:** {tag_badge(row.get('ai_tag', '其他'))}", unsafe_allow_html=True)
                st.markdown(f"**信号判断:** {signal_badge(row.get('ai_signal', '中性'))}", unsafe_allow_html=True)
                if row.get("url"):
                    st.link_button("查看原文 🔗", row["url"], width="stretch")
                if row.get("pdf_url"):
                    st.link_button("下载 PDF 📄", row["pdf_url"], width="stretch")

                # 生成结构化深度分析并提供单条下载
                safe_title = re.sub(r'[\\/:*?"<>|]', "_", row['title'])[:40]
                report_filename = f"{row['announce_date']}_{stock_code}_{safe_title}_AI深度分析.docx"
                report_docx = generate_report_docx(stock_code, name, row)
                st.download_button(
                    label="📝 下载深度分析报告",
                    data=report_docx,
                    file_name=report_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"dl_{stock_code}_{row.get('art_code', idx)}",
                    width="stretch",
                )

            with cols[1]:
                summary = get_structured_summary(row, stock_name=name)
                st.markdown("**📋 结构化摘要**")
                st.info(summary)

                reason = get_reason(row)
                st.markdown("**判断理由**")
                st.success(reason)

    st.caption("提示：结构化摘要和判断理由优先展示本地数据库中的 Kimi AI 分析结果；若无，则尝试调用 LLM API（需配置 MOONSHOT_API_KEY 或 OPENAI_API_KEY）；否则使用规则模板生成。")
